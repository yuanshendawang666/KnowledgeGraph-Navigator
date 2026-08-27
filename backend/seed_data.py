"""
样本数据生成脚本
----------------
自动创建教师、学生、课程，上传文档，
触发DeepSeek知识提取，设置学习进度，
并测试推荐与问答功能。
"""

import os
import sys
import time
import requests

BASE = "http://localhost:8000"

# ── 辅助函数 ────────────────────────────────────
def api(method, path, **kwargs):
    """统一的 API 调用，自动处理 auth header"""
    headers = kwargs.pop("headers", {})
    if hasattr(api, "token"):
        headers["Authorization"] = f"Bearer {api.token}"
    kwargs.setdefault("timeout", 120)
    if method == "GET":
        return requests.get(f"{BASE}{path}", headers=headers, **kwargs)
    return requests.post(f"{BASE}{path}", headers=headers, **kwargs)

def ok(r, label=""):
    icon = "[OK]" if r.status_code < 400 else "[FAIL]"
    print(f"  {icon} {label} [{r.status_code}]")
    return r.json() if r.status_code < 400 else None

# ── Step 1: 创建用户 ─────────────────────────────
print("\n" + "=" * 60)
print("  Step 1: 创建用户")
print("=" * 60)

def ensure_user(username, email, password, role):
    r = requests.post(f"{BASE}/api/auth/register", json={
        "username": username, "email": email,
        "password": password, "role": role,
    })
    if r.status_code == 201:
        ok(r, f"注册{role}: {username}")
        return r.json()["access_token"]
    else:
        print(f"  [INFO] {username} 已存在，直接登录")
        r = requests.post(f"{BASE}/api/auth/login", json={
            "username": username, "password": password,
        })
        return ok(r, f"登录: {username}")["access_token"]

api.token = ensure_user("teacher_wang", "wang@school.edu.cn", "123456", "teacher")
ensure_user("student_ming", "ming@school.edu.cn", "123456", "student")

# 确保用的是 teacher 的 token（后续操作需要教师权限）
r = requests.post(f"{BASE}/api/auth/login", json={
    "username": "teacher_wang", "password": "123456"
})
api.token = r.json()["access_token"]

# ── Step 2: 创建课程 ─────────────────────────────
print("\n" + "=" * 60)
print("  Step 2: 创建课程")
print("=" * 60)

# 课程 1: Python 程序设计
r = api("POST", "/api/courses/", json={
    "title": "Python程序设计基础",
    "description": "面向大一新生的Python入门课程，涵盖变量、控制流、函数、面向对象等核心概念",
})
course1 = ok(r, "创建课程: Python程序设计基础")

# 课程 2: 数据结构
r = api("POST", "/api/courses/", json={
    "title": "数据结构与算法",
    "description": "计算机科学核心课程，线性表、栈队列、树、图、排序算法",
})
course2 = ok(r, "创建课程: 数据结构与算法")

cid1 = course1["id"]
cid2 = course2["id"]

# ── Step 3: 上传文档 & 知识提取 ──────────────────
print("\n" + "=" * 60)
print("  Step 3: 上传文档 → DeepSeek 知识提取")
print("=" * 60)

def create_and_upload(course_id, filename, content):
    """创建docx文档并上传"""
    from docx import Document
    import tempfile
    doc = Document()
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], 1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], 2)
        else:
            doc.add_paragraph(line)
    tmp = os.path.join(tempfile.gettempdir(), filename)
    doc.save(tmp)
    with open(tmp, "rb") as f:
        r = api("POST", f"/api/courses/{course_id}/upload",
                files={"file": (filename, f,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    return ok(r, f"上传 {filename}")

def extract_knowledge(course_id, label):
    r = api("POST", f"/api/courses/{course_id}/extract", timeout=180)
    data = ok(r, f"DeepSeek提取: {label}")
    if data:
        kps = data.get("knowledge_points", [])
        rels = data.get("relations", [])
        print(f"        {len(kps)}个知识点, {len(rels)}个关系")
        for kp in kps:
            print(f"        · {kp['name']}")
        for rel in rels[:8]:
            print(f"          {rel['source']} --[{rel['relation_type']}]--> {rel['target']}")
    return data

# 课程1: Python — 树状结构内容
py_content = """# Python程序设计基础

# 第一章：基础语法
## 变量与数据类型
变量是存储数据的容器。Python是动态类型语言，支持整数int、浮点数float、字符串str、布尔值bool等基础类型。
整数用于表示整数值，如年龄、计数。浮点数用于小数计算。字符串用于文本处理。布尔值表示真和假。

## 运算符与表达式
算术运算符(加减乘除)、比较运算符(大于小于等于)、逻辑运算符(and/or/not)是编写程序逻辑的基础。运算符优先级决定表达式的计算顺序。

## 输入与输出
print函数用于输出信息到控制台，支持格式化输出。input函数接收用户键盘输入，返回字符串类型。类型转换函数int()、float()、str()用于数据类型互转。

# 第二章：控制流程
## 条件判断
if语句根据条件执行不同代码块。Python用缩进定义代码块。完整结构包括if、elif和else。条件可以是比较表达式或逻辑表达式。

## 循环结构
for循环遍历可迭代对象(列表、字符串、range)。while循环在条件为真时反复执行。break跳出循环，continue跳过当前迭代。循环可以嵌套使用。

# 第三章：函数与模块
## 函数定义
def关键字定义函数。函数封装可重用代码，接收参数并返回值。支持位置参数、默认参数、可变参数和关键字参数。函数提高代码的模块化程度。

## 内置函数与模块
Python提供丰富的内置函数(len/type/range等)和标准库模块。import导入模块，from...import导入特定函数。os、math、random等常用模块。

# 第四章：数据结构
## 列表操作
列表是有序可变的序列。支持索引访问、切片、append/extend/insert等操作。列表推导式提供简洁的创建方式。列表适合存储有序集合。

## 字典与集合
字典是键值对的无序集合，通过键快速访问值。集合是无重复元素的集合，支持交并差运算。字典适合键值映射场景。

## 字符串处理
字符串是不可变序列。支持切片、拼接、格式化(f-string)。常用方法：split分割、join连接、replace替换、strip去空格。

# 第五章：面向对象
## 类与对象
类class定义对象的属性和方法。__init__是构造方法，self代表实例本身。对象是类的实例。类是面向对象编程的核心概念。

## 继承与多态
继承允许子类复用父类代码，支持单继承和多继承。多态使不同类对象通过统一接口调用。方法重写允许子类定制父类行为。
"""

create_and_upload(cid1, "python_basics.docx", py_content)
extract_knowledge(cid1, "Python程序设计基础")

# 课程2: 数据结构
ds_content = """# 数据结构与算法

## 算法复杂度分析
时间复杂度描述算法执行时间随输入规模增长的趋势，使用大O表示法。空间复杂度描述内存占用。常见复杂度：O(1)、O(log n)、O(n)、O(n log n)、O(n²)。

## 线性表
线性表是最基本的数据结构，元素之间是一对一的关系。顺序表用数组实现，支持随机访问。链表用指针连接节点，插入删除高效。

## 栈与队列
栈是后进先出LIFO的结构，push入栈、pop出栈。队列是先进先出FIFO的结构，enqueue入队、dequeue出队。它们都可以用顺序表或链表实现。

## 树与二叉树
树是层次化的数据结构。二叉树每个节点最多两个子节点。二叉搜索树左小右大，支持快速查找。遍历方式有前序、中序、后序和层序。

## 排序算法
冒泡排序每次将最大的元素冒到末尾，O(n²)。快速排序用分治法选基准分区，平均O(n log n)。归并排序先分后合，稳定O(n log n)。
"""

create_and_upload(cid2, "data_structure.docx", ds_content)
extract_knowledge(cid2, "数据结构与算法")

# ── Step 4: 学习进度 ─────────────────────────────
print("\n" + "=" * 60)
print("  Step 4: 学习进度 (学生 ming)")
print("=" * 60)

# 切换到学生账号
r = requests.post(f"{BASE}/api/auth/login", json={
    "username": "student_ming", "password": "123456"
})
api.token = r.json()["access_token"]

# 获取课程1的知识点列表
from app.core.database import SessionLocal
from app.models import KnowledgePoint

db = SessionLocal()
kps1 = db.query(KnowledgePoint).filter(KnowledgePoint.course_id == cid1).order_by(KnowledgePoint.order_index).all()
kps2 = db.query(KnowledgePoint).filter(KnowledgePoint.course_id == cid2).order_by(KnowledgePoint.order_index).all()
db.close()

# 课程1: 前2个已掌握，第3个学习中
for kp in kps1[:2]:
    api("POST", "/api/learning/progress", json={
        "knowledge_point_id": kp.id, "status": "mastered", "mastery_level": 0.95
    })
if len(kps1) > 2:
    api("POST", "/api/learning/progress", json={
        "knowledge_point_id": kps1[2].id, "status": "in_progress", "mastery_level": 0.5
    })

# 课程2: 第1个已掌握
if kps2:
    api("POST", "/api/learning/progress", json={
        "knowledge_point_id": kps2[0].id, "status": "mastered", "mastery_level": 0.9
    })

# 统计
for cid, label in [(cid1, "Python程序设计"), (cid2, "数据结构")]:
    r = api("GET", f"/api/learning/stats/{cid}")
    s = ok(r, f"进度统计: {label}")
    if s:
        print(f"        总{s['total_points']}个 | 已掌握{s['mastered_count']} | 学习中{s['in_progress_count']} | 进度{s['progress_percentage']}%")

# ── Step 5: 学习路径推荐 ─────────────────────────
print("\n" + "=" * 60)
print("  Step 5: 个性化路径推荐")
print("=" * 60)

for cid, label in [(cid1, "Python程序设计"), (cid2, "数据结构")]:
    r = api("GET", f"/api/learning/recommend/{cid}")
    rec = ok(r, f"推荐路径: {label}")
    if rec:
        ready = rec.get("ready_to_learn", [])
        next_kps = rec.get("recommended_next", [])
        print(f"        总知识点{rec['total_count']} | 已掌握{rec['mastered_count']}")
        print(f"        可学{len(ready)}个 | 推荐{len(next_kps)}个")
        for kp in next_kps:
            print(f"        → {kp['label']}")

# ── Step 6: 图谱数据 ─────────────────────────────
print("\n" + "=" * 60)
print("  Step 6: 知识图谱")
print("=" * 60)

for cid, label in [(cid1, "Python程序设计"), (cid2, "数据结构")]:
    r = api("GET", f"/api/courses/{cid}/graph")
    g = ok(r, f"图谱: {label}")
    if g:
        print(f"        {len(g['nodes'])}节点 {len(g['edges'])}边 — 可直接用于AntV G6渲染")

# ── Step 7: 智能问答 ─────────────────────────────
print("\n" + "=" * 60)
print("  Step 7: 智能问答 (RAG)")
print("=" * 60)

questions = [
    (cid1, "Python中列表和字典有什么区别？"),
    (cid1, "什么是面向对象的三大特性？"),
    (cid2, "快速排序的时间复杂度是多少？为什么？"),
]

for cid, q in questions:
    r = api("POST", "/api/qa/ask", json={"course_id": cid, "question": q}, timeout=60)
    data = ok(r, f"Q: {q[:30]}...")
    if data and data.get("code") == 0:
        ans = data["data"]["answer"][:120].replace("\n", " ")
        refs = [ref["name"] for ref in data["data"].get("references", [])]
        print(f"        A: {ans}...")
        if refs:
            print(f"        参考: {refs}")

# 推荐问题
r = api("GET", "/api/qa/recommend-questions?course_id=1", timeout=60)
data = ok(r, "推荐问题")
if data and data.get("code") == 0:
    questions = data["data"].get("questions", [])
    for q in questions:
        print(f"        - {q}")

# ── 总结 ────────────────────────────────────────
print("\n" + "=" * 60)
print("  样本数据生成完毕！")
print("=" * 60)
print(f"""
  教师: teacher_wang / 123456
  学生: student_ming / 123456

  课程1 (id={cid1}): Python程序设计基础
  课程2 (id={cid2}): 数据结构与算法

  Swagger: http://localhost:8000/docs
  API:     http://localhost:8000/
""")
