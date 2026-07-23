"""
文档解析服务
------------
支持 PDF 和 DOCX 格式文档的文本提取与中文预处理。

技术栈：
- pdfplumber: 解析 PDF 文件
- python-docx: 解析 DOCX 文件
- jieba: 中文分词与关键词提取
- re: 正则表达式文本清洗

使用方式：
    from app.services.parser import DocumentParser

    parser = DocumentParser()
    text = parser.parse("path/to/file.pdf")
    keywords = parser.extract_keywords(text, topk=20)
"""

import os
import re
from collections import Counter
from pathlib import Path

import pdfplumber
import docx
import jieba
import jieba.analyse


class DocumentParser:
    """文档解析器，支持 PDF 和 DOCX 格式"""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}

    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """检查文件格式是否受支持"""
        ext = Path(filename).suffix.lower()
        return ext in cls.SUPPORTED_EXTENSIONS

    def parse(self, file_path: str) -> str:
        """
        解析文档，返回提取的纯文本内容。

        Args:
            file_path: 文档文件的绝对路径

        Returns:
            提取的文本内容

        Raises:
            ValueError: 不支持的文件格式
            FileNotFoundError: 文件不存在
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            return self._parse_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    def _parse_pdf(self, file_path: str) -> str:
        """解析 PDF 文件，逐页提取文本"""
        text_parts = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        # 清理多余空白
                        cleaned = self._clean_text(page_text)
                        text_parts.append(cleaned)
        except Exception as e:
            raise RuntimeError(f"PDF 解析失败: {e}")

        return "\n\n".join(text_parts)

    def _parse_docx(self, file_path: str) -> str:
        """解析 DOCX 文件，逐段提取文本"""
        text_parts = []
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    cleaned = self._clean_text(para.text)
                    text_parts.append(cleaned)

            # 也提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)
        except Exception as e:
            raise RuntimeError(f"DOCX 解析失败: {e}")

        return "\n\n".join(text_parts)

    def _clean_text(self, text: str) -> str:
        """清理文本：去除多余空白、统一换行"""
        # 将多个空格合并为一个
        text = re.sub(r"[ \t]+", " ", text)
        # 将多个换行合并为最多两个
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    # ---- 中文文本预处理 ----

    @staticmethod
    def segment(text: str) -> list:
        """
        使用 jieba 对中文文本进行分词。

        Args:
            text: 待分词的文本

        Returns:
            分词后的词语列表
        """
        return list(jieba.cut(text))

    @staticmethod
    def extract_keywords(text: str, topk: int = 20) -> list:
        """
        使用 jieba TF-IDF 算法提取文本关键词。

        Args:
            text: 待提取的文本
            topk: 返回的关键词数量

        Returns:
            [(关键词, 权重), ...] 按权重降序排列
        """
        return jieba.analyse.extract_tags(text, topK=topk, withWeight=True)

    @staticmethod
    def extract_keywords_textrank(text: str, topk: int = 20) -> list:
        """
        使用 jieba TextRank 算法提取文本关键词。

        Args:
            text: 待提取的文本
            topk: 返回的关键词数量

        Returns:
            [(关键词, 权重), ...] 按权重降序排列
        """
        return jieba.analyse.textrank(text, topK=topk, withWeight=True)

    @staticmethod
    def get_word_frequency(text: str, min_len: int = 2) -> list:
        """
        统计词频（过滤单字词和标点）。

        Args:
            text: 待统计的文本
            min_len: 最小词长

        Returns:
            [(词语, 频次), ...] 按频次降序排列
        """
        words = jieba.cut(text)
        counter = Counter(
            w for w in words if len(w.strip()) >= min_len
        )
        return counter.most_common(100)


# 便捷函数
def parse_document(file_path: str) -> str:
    """解析文档的便捷函数"""
    parser = DocumentParser()
    return parser.parse(file_path)


def extract_keywords(text: str, topk: int = 20) -> list:
    """提取关键词的便捷函数"""
    return DocumentParser.extract_keywords(text, topk)

